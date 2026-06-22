from docx import Document
import os

def generate_docx(content, filename):

    os.makedirs("exports", exist_ok=True)

    doc = Document()

    doc.add_heading(
        "Software Requirements Specification",
        level=1
    )

    doc.add_paragraph(content)

    filepath = f"exports/{filename}.docx"

    doc.save(filepath)

    return filepath